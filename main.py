<<<<<<< HEAD
from flask import Flask, request, jsonify
from flasgger import Swagger
import PyPDF2
import docx
from pinecone import Pinecone
from FlagEmbedding import BGEM3FlagModel
import uuid


app = Flask(__name__)
swagger = Swagger(app)


model = BGEM3FlagModel('BAAI/bge-m3', use_fp16=True)


pc = Pinecone(api_key="pcsk_3CPrZH_65UmbWfwGXiJYiFqzciZudoXx57Au2F1jNysFMvTAj5tZqbQriUtR7o8wRSGTda")
index = pc.Index("thuctap") 


def extract_pdf(file_stream):
    reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_docx(file_stream):
    doc = docx.Document(file_stream)
    return "\n".join([p.text for p in doc.paragraphs])


# ================= CHUNKING =================
def chunk_text(text, max_len=300):
    sentences = text.split('\n')
    chunks = []
    current = ""

    for s in sentences:
        if len(current) + len(s) < max_len:
            current += " " + s
        else:
            chunks.append(current.strip())
            current = s

    if current:
        chunks.append(current.strip())

    return chunks


# ================= EMBEDDING =================
def embed_texts(texts):
    embeddings = model.encode(texts)['dense_vecs']
    return [emb.tolist() for emb in embeddings]


# ================= API: UPLOAD CV =================
@app.route('/upload-cv', methods=['POST'])
def upload_cv():
    """
    Upload CV and store embeddings
    ---
    consumes:
      - multipart/form-data
    parameters:
      - name: file
        in: formData
        type: file
        required: true
    responses:
      200:
        description: CV processed
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    filename = file.filename.lower()

    # Extract text
    if filename.endswith('.pdf'):
        text = extract_pdf(file)
    elif filename.endswith('.docx'):
        text = extract_docx(file)
    else:
        return jsonify({"error": "Unsupported format"}), 400

    # Chunk
    chunks = chunk_text(text)

    # Embed
    embeddings = embed_texts(chunks)

    # Store
    vectors = []
    base_id = str(uuid.uuid4())

    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        vectors.append({
            "id": f"{base_id}_{i}",
            "values": emb,
            "metadata": {
                "type": "cv",
                "filename": file.filename,
                "text": chunk
            }
        })

    index.upsert(vectors)

    return jsonify({
        "message": "CV stored successfully",
        "chunks": len(chunks)
    })


# ================= API: ADD JOB =================
@app.route('/add-job', methods=['POST'])
def add_job():
    """
    Add Job Description
    ---
    parameters:
      - name: body
        in: body
        required: true
        schema:
          properties:
            text:
              type: string
    """
    data = request.json
    job_text = data.get("text")

    if not job_text:
        return jsonify({"error": "No job text"}), 400

    chunks = chunk_text(job_text)
    embeddings = embed_texts(chunks)

    vectors = []
    base_id = str(uuid.uuid4())

    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        vectors.append({
            "id": f"{base_id}_{i}",
            "values": emb,
            "metadata": {
                "type": "job",
                "text": chunk
            }
        })

    index.upsert(vectors)

    return jsonify({
        "message": "Job stored successfully",
        "chunks": len(chunks)
    })


# ================= API: MATCH =================
@app.route('/match', methods=['POST'])
def match():
    """
    Match Job to CVs
    ---
    parameters:
      - name: body
        in: body
        required: true
        schema:
          properties:
            text:
              type: string
    """
    data = request.json
    job_text = data.get("text")

    if not job_text:
        return jsonify({"error": "No job text"}), 400

    # Embed job query
    query_embedding = embed_texts([job_text])[0]

    # Search
    results = index.query(
        vector=query_embedding,
        top_k=5,
        include_metadata=True,
        filter={"type": {"$eq": "cv"}}
    )

    matches = []
    for match in results['matches']:
        matches.append({
            "score": match['score'],
            "filename": match['metadata']['filename'],
            "text": match['metadata']['text']
        })

    return jsonify({
        "results": matches
    })


if __name__ == "__main__":
=======
from flask import Flask, request, jsonify
from flasgger import Swagger
import PyPDF2
import docx
from pinecone import Pinecone
from FlagEmbedding import BGEM3FlagModel
import uuid


app = Flask(__name__)
swagger = Swagger(app)


model = BGEM3FlagModel('BAAI/bge-m3', use_fp16=True)


pc = Pinecone(api_key="pcsk_3CPrZH_65UmbWfwGXiJYiFqzciZudoXx57Au2F1jNysFMvTAj5tZqbQriUtR7o8wRSGTda")
index = pc.Index("thuctap") 


def extract_pdf(file_stream):
    reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def extract_docx(file_stream):
    doc = docx.Document(file_stream)
    return "\n".join([p.text for p in doc.paragraphs])


# ================= CHUNKING =================
def chunk_text(text, max_len=300):
    sentences = text.split('\n')
    chunks = []
    current = ""

    for s in sentences:
        if len(current) + len(s) < max_len:
            current += " " + s
        else:
            chunks.append(current.strip())
            current = s

    if current:
        chunks.append(current.strip())

    return chunks


# ================= EMBEDDING =================
def embed_texts(texts):
    embeddings = model.encode(texts)['dense_vecs']
    return [emb.tolist() for emb in embeddings]


# ================= API: UPLOAD CV =================
@app.route('/upload-cv', methods=['POST'])
def upload_cv():
    """
    Upload CV and store embeddings
    ---
    consumes:
      - multipart/form-data
    parameters:
      - name: file
        in: formData
        type: file
        required: true
    responses:
      200:
        description: CV processed
    """
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    filename = file.filename.lower()

    # Extract text
    if filename.endswith('.pdf'):
        text = extract_pdf(file)
    elif filename.endswith('.docx'):
        text = extract_docx(file)
    else:
        return jsonify({"error": "Unsupported format"}), 400

    # Chunk
    chunks = chunk_text(text)

    # Embed
    embeddings = embed_texts(chunks)

    # Store
    vectors = []
    base_id = str(uuid.uuid4())

    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        vectors.append({
            "id": f"{base_id}_{i}",
            "values": emb,
            "metadata": {
                "type": "cv",
                "filename": file.filename,
                "text": chunk
            }
        })

    index.upsert(vectors)

    return jsonify({
        "message": "CV stored successfully",
        "chunks": len(chunks)
    })


# ================= API: ADD JOB =================
@app.route('/add-job', methods=['POST'])
def add_job():
    """
    Add Job Description
    ---
    parameters:
      - name: body
        in: body
        required: true
        schema:
          properties:
            text:
              type: string
    """
    data = request.json
    job_text = data.get("text")

    if not job_text:
        return jsonify({"error": "No job text"}), 400

    chunks = chunk_text(job_text)
    embeddings = embed_texts(chunks)

    vectors = []
    base_id = str(uuid.uuid4())

    for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
        vectors.append({
            "id": f"{base_id}_{i}",
            "values": emb,
            "metadata": {
                "type": "job",
                "text": chunk
            }
        })

    index.upsert(vectors)

    return jsonify({
        "message": "Job stored successfully",
        "chunks": len(chunks)
    })


# ================= API: MATCH =================
@app.route('/match', methods=['POST'])
def match():
    """
    Match Job to CVs
    ---
    parameters:
      - name: body
        in: body
        required: true
        schema:
          properties:
            text:
              type: string
    """
    data = request.json
    job_text = data.get("text")

    if not job_text:
        return jsonify({"error": "No job text"}), 400

    # Embed job query
    query_embedding = embed_texts([job_text])[0]

    # Search
    results = index.query(
        vector=query_embedding,
        top_k=5,
        include_metadata=True,
        filter={"type": {"$eq": "cv"}}
    )

    matches = []
    for match in results['matches']:
        matches.append({
            "score": match['score'],
            "filename": match['metadata']['filename'],
            "text": match['metadata']['text']
        })

    return jsonify({
        "results": matches
    })


if __name__ == "__main__":
>>>>>>> 2889c005460580cfc56524371d405006df9a3b8a
    app.run(debug=True)